from docx import Document


class Docxfile:
    # init a new docx

    def __init__(self, path):
        self.path = path
        # self.name = name
        # self.doc = Document()

    # import a exist docx
    def genDocument(self):
        # self.path = path
        self.doc = Document(self.path)


    def replaceContent(self, key, sub):
        # parag = 0
        for para in self.doc.paragraphs:
            if key in para.text:
                inline = para.runs
                print('SEARCH FOUND!!')
                for ik in range(len(inline)):
                    if key in inline[ik].text:
                        text = inline[ik].text.replace(key, sub)
                        inline[ik].text = text

            # parag+=1
        # self.doc.save('test.docx')
        # self.doc = Document('test.docx')
        return 0

    def saveFile(self, filename):
        if self.doc.save(filename+'.docx'):
            return 0
        else:
            return -1
